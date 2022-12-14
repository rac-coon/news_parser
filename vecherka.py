# -*- coding: utf-8 -*-
import requests
import docx
import re
import time
import os
from bs4 import BeautifulSoup

if __name__ == "__main__":

    __start_time__ = time.time()

    # Получение исходных статей

    #req_main = requests.get("https://vm.ru/society")
    req_main = requests.get("https://vm.ru/politics")
    src_main = req_main.text

    # with open("politica1.html", encoding="utf-8") as file:
    #     src_main = file.read()
    #
    soup_main = BeautifulSoup(src_main, "lxml")

    # Сбор ссылок

    articles_lib = soup_main.select('.listing.category-listing > .topic.litem.article')
    urls_lib = []
    for article in articles_lib:
        #urls_lib.append(article.find(class_='wst').get("href"))
        urls_lib.append("https://vm.ru" + article.find(class_='wst').get("href"))

    # Начальные значения для номеров файлов

    s_count = 175
    m_count = 175
    l_count = 175

    # Проверка количества отработанных ссылок

    print(len(urls_lib), "ссылок")
    __test_count__ = 0

    # Счётчик получившихся текстов

    __s_articles_count__ = 0
    __m_articles_count__ = 0
    __l_articles_count__ = 0

    # Создание рабочих каталогов

    if not os.path.exists("S"):
        os.makedirs("S/")
        print("Была создана директория S")
    if not os.path.exists("M"):
        os.makedirs("M/")
        print("Была создана директория M")
    if not os.path.exists("L"):
        print("Была создана директория L")
        os.makedirs("L/")

    # Работа со статьями

    for url in urls_lib:
        req = requests.get(url)

        soup = BeautifulSoup(req.text, "lxml")

    # Переменные для подсчёта количества и записи в файл

        words_count = 0
        chars_count = 0
        all_chars_count = 0
        article_text = ""

    # Счётчик отработанных ссылок

        __test_count__ += 1

    # Сбор текста

        for tab in soup.select('.article-view > p'):
            article_text += str(tab.text + "\n")
            tab_text = str(tab.text)
            tab_text.replace('-', '')
            words_count += len(re.findall(r'\w+', tab_text))
            all_chars_count += len(tab_text)
            chars_count += len(re.findall(r'\S', tab_text))
            if tab.findAll(text="На правах рекламы"):
                print("реклама")
                break

    # Запись текста в необходимый .docx и запись метаданных .txt

        if (words_count > 50) and (words_count < 250):
            __s_articles_count__ += 1
            if s_count > 209:
                continue
        # больше не нада
            continue
            with open("S/" + str(s_count) + ".txt", "w", encoding="utf-8") as file:
                file.write(f"{url}\n13.04.2022\n{words_count}\n{all_chars_count}, {chars_count}\n7\n")
            new_doc = docx.Document()
            new_doc.add_paragraph(article_text)
            new_doc.save("E:/PycharmProjects/S/" + str(s_count) + ".docx")
            s_count += 1
        if (words_count > 251) and (words_count < 900):
            __m_articles_count__ += 1
            if m_count > 209:
                continue
        # больше не нада
            continue
            with open("M/" + str(m_count) + ".txt", "w", encoding="utf-8") as file:
                file.write(f"{url}\n07.04.2022\n{words_count}\n{all_chars_count}, {chars_count}\n7\n")
            new_doc = docx.Document()
            new_doc.add_paragraph(article_text)
            new_doc.save("E:/PycharmProjects/M/" + str(m_count) + ".docx")
            m_count += 1
        if (words_count > 901) and (words_count < 2000):
            __l_articles_count__ += 1
            if l_count > 209:
                continue
            with open("L/" + str(l_count) + ".txt", "w", encoding="utf-8") as file:
                file.write(f"{url}\n07.04.2022\n{words_count}\n{all_chars_count}, {chars_count}\n7\n")
            new_doc = docx.Document()
            new_doc.add_paragraph(article_text)
            new_doc.save("E:/PycharmProjects/L/" + str(l_count) + ".docx")
            l_count += 1

    # Итог работы программы: количество отработанных ссылок и время работы

    print(f"{__test_count__} ссылок обработано\n"
          f"S: {__s_articles_count__}/34\nM: {__m_articles_count__}/34\nL: {__l_articles_count__}/34\n"
          f"Время работы: {time.time() - __start_time__}")
